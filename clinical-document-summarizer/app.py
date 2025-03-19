import streamlit as st
import anthropic
import google.generativeai as genai
import os
import fitz  # PyMuPDF
import json
from docx import Document
from dotenv import load_dotenv
from fpdf import FPDF
import tempfile

# Load API keys from environment variables or secrets
def get_api_keys():
    """Get API keys from environment variables or Streamlit secrets"""
    try:
        anthropic_key = st.secrets["ANTHROPIC_API_KEY"]
        gemini_key = st.secrets["GEMINI_API_KEY"]
    except:
        load_dotenv()
        anthropic_key = os.getenv("ANTHROPIC_API_KEY", "")
        gemini_key = os.getenv("GEMINI_API_KEY", "")
    return anthropic_key, gemini_key

# Initialize AI clients
def initialize_clients(anthropic_key, gemini_key):
    """Initialize both Anthropic and Gemini clients"""
    try:
        anthropic_client = anthropic.Client(api_key=anthropic_key) if anthropic_key else None
        if gemini_key:
            genai.configure(api_key=gemini_key)
            gemini_client = genai.GenerativeModel("gemini-1.5-pro")
        else:
            gemini_client = None
        return anthropic_client, gemini_client
    except Exception as e:
        st.error(f"Error initializing clients: {str(e)}")
        return None, None

# Extract text from PDF using PyMuPDF
def extract_text_from_pdf(uploaded_file):
    """Extract text from an uploaded PDF file"""
    text = ""
    try:
        # Create a temporary file to store the uploaded content
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            # Write the uploaded file content to the temporary file
            tmp_file.write(uploaded_file.getvalue())
            tmp_file.flush()
            
            # Open the temporary file with PyMuPDF
            with fitz.open(tmp_file.name) as doc:
                for page in doc:
                    text += page.get_text("text") + "\n"
            
            # Delete the temporary file
            os.unlink(tmp_file.name)
            
        return text.strip()
    except Exception as e:
        st.error(f"Error extracting text from PDF: {str(e)}")
        return ""

# Use Gemini 1.5 Pro to parse the full document
def parse_document_with_gemini(text, gemini_client):
    prompt = f"""
    Analyze this clinical protocol document and extract its complete content.
    Maintain all the important details, context, and relationships between different sections.
    Format the response as detailed text that preserves the document's information and structure.
    
    Document:
    {text}
    """
    try:
        response = gemini_client.generate_content(prompt)
        return response.text, None
    except Exception as e:
        return None, str(e)

# Use Claude to parse the full document
def parse_document_with_claude(text, claude_client):
    prompt = f"""
    Analyze this clinical protocol document and extract its complete content.
    Maintain all the important details, context, and relationships between different sections.
    Format the response as detailed text that preserves the document's information and structure.
    
    Document:
    {text}
    """
    try:
        response = claude_client.messages.create(
            model="claude-3-5-haiku-20241022",
            max_tokens=4000,
            temperature=0.1,
            system="You are a clinical protocol analyst. Extract and structure the complete document content.",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.content[0].text.strip(), None
    except Exception as e:
        return None, str(e)

# Parse document with fallback
def parse_document(text, gemini_client, claude_client):
    """Parse document using Gemini first, falling back to Claude if Gemini fails"""
    
    # Try Gemini first
    if gemini_client:
        st.info("🔄 Attempting to parse with Gemini...")
        content, error = parse_document_with_gemini(text, gemini_client)
        if content:
            st.success("✅ Document parsed successfully with Gemini")
            return content
        else:
            st.warning(f"⚠️ Gemini parsing failed: {error}")
    
    # Fall back to Claude
    if claude_client:
        st.info("🔄 Falling back to Claude for parsing...")
        content, error = parse_document_with_claude(text, claude_client)
        if content:
            st.success("✅ Document parsed successfully with Claude")
            return content
        else:
            st.error(f"❌ Claude parsing failed: {error}")
    
    # If both fail, return original text
    st.error("❌ Both parsing attempts failed. Using original text.")
    return text

# Use Claude to extract and format specific fields
def extract_field_with_claude(parsed_content, field_title, field_description, anthropic_client):
    prompt = f"""
    You are analyzing a clinical study protocol. For the field specified below, extract and summarize the most relevant information from the parsed document.
    
    Field Title: {field_title}
    Expected Content: {field_description}
    
    Guidelines:
    1. Focus specifically on content relevant to this field
    2. Provide comprehensive but concise information
    3. Maintain professional medical terminology
    4. Include all relevant details and context
    5. Format the response in clear, readable paragraphs
    
    Parsed Document Content:
    {parsed_content}
    """
    
    try:
        response = anthropic_client.messages.create(
            model="claude-3-5-haiku-20241022",
            max_tokens=1000,
            temperature=0.1,
            system="You are a clinical protocol analyst. Extract and format relevant information for the specified field.",
            messages=[{"role": "user", "content": prompt}]
        )
        content = response.content[0].text.strip()
        return content if content else "No relevant information found for this field."
    except Exception as e:
        st.error(f"Error extracting field with Claude: {str(e)}")
        return f"Error processing {field_title}: {str(e)}"

# Generate DOCX summary
def generate_docx_summary(extracted_fields, output_path="summary.docx"):
    doc = Document()
    doc.add_heading("Clinical Study Summary", level=1)
    
    for field_title, content in extracted_fields.items():
        doc.add_heading(field_title, level=2)
        doc.add_paragraph(content)
    
    doc.save(output_path)
    return output_path

# Main execution
def main():
    st.set_page_config(page_title="Clinical Protocol Summary Generator", layout="wide")
    
    st.title("Clinical Protocol Summary Generator")
    st.write("Upload a clinical study protocol document to generate a summarized version")

    # Configuration in sidebar
    with st.sidebar:
        st.header("Configuration")
        
        # Get API keys
        anthropic_key, gemini_key = get_api_keys()
        
        # API Key input fields
        new_anthropic_key = st.text_input(
            "Anthropic API Key", 
            type="password",
            value=anthropic_key,
            help="Enter your Anthropic API key"
        )
        
        new_gemini_key = st.text_input(
            "Gemini API Key",
            type="password",
            value=gemini_key,
            help="Enter your Gemini API key"
        )
        
        # Update API keys if changed
        anthropic_key = new_anthropic_key if new_anthropic_key else anthropic_key
        gemini_key = new_gemini_key if new_gemini_key else gemini_key
        
        # Initialize clients
        anthropic_client, gemini_client = initialize_clients(anthropic_key, gemini_key)
        
        # Check if at least one API key is available
        if not anthropic_key and not gemini_key:
            st.error("❌ At least one API key is required")
            st.info("💡 Get your API keys from:")
            st.info("• Anthropic: [Console](https://console.anthropic.com)")
            st.info("• Google: [AI Studio](https://makersuite.google.com/app/apikey)")
            st.stop()
        elif not anthropic_client and not gemini_client:
            st.error("❌ Error initializing AI clients")
            st.stop()
        else:
            available_services = []
            if anthropic_client:
                available_services.append("Claude")
            if gemini_client:
                available_services.append("Gemini")
            st.success(f"✅ Available services: {', '.join(available_services)}")
        
        # File uploader
        uploaded_file = st.file_uploader("Upload Clinical Protocol PDF", type=["pdf"])
    
    if uploaded_file and st.button("Process Document"):
        try:
            with st.spinner("📄 Extracting text from PDF..."):
                pdf_text = extract_text_from_pdf(uploaded_file)
                if not pdf_text:
                    st.error("❌ Could not extract text from PDF")
                    st.stop()
            
            # Parse document with fallback
            parsed_content = parse_document(pdf_text, gemini_client, anthropic_client)
            if not parsed_content:
                st.error("❌ Could not parse document with any available service")
                st.stop()
            
            # Define fields to extract
            field_definitions = {
                "STUDY TITLE": "Provide full title of the study",
                "CLINICAL PHASE": "Specify clinical phase (1, 2a)",
                "STUDY OBJECTIVES": "Provide a brief description of the study objectives, including primary, secondary, and exploratory objectives.",
                "STUDY RATIONALE": "Summarize the rationale for testing the proposed therapy.",
                "STUDY POPULATION": "Briefly describe the study population and explain the rationale for choosing this population.",
                "MAIN INCLUSION/EXCLUSION CRITERIA": "Specify the main inclusion/exclusion criteria and explain the rationale.",
                "PRIMARY ENDPOINT(S)": "Describe the Primary Endpoint(s) and the set of measurements used to address the objectives.",
                "SECONDARY & EXPLORATORY ENDPOINTS": "Describe the Secondary & Exploratory Endpoint(s) and measures that will address them.",
                "STUDY DESIGN": "Summarize the study design, including type of study, number of arms, controls or comparators.",
                "SUBJECT NUMBER": "Provide the total number of study subjects, the number per study arm, and justification.",
                "TREATMENT DURATION": "Specify the length of the treatment period.",
                "DURATION OF FOLLOW UP": "Specify the length of the protocol-specified follow-up period.",
                "DOSE LEVEL(S) AND DOSE JUSTIFICATION": "Specify the dose level(s), number of doses, and dosing frequency. Summarize how dosing was determined.",
                "ROUTE OF DELIVERY": "Specify how the doses will be delivered.",
                "DATA and SAFETY MONITORING PLAN": "Summarize the Data and Safety Monitoring Plan.",
                "STOPPING RULES": "Specify stopping rules.",
                "IMMUNE MONITORING & IMMUNOSUPPRESSION": "Describe and justify the plan for immunosuppression and immune monitoring (if applicable).",
                "SUPPORTING STUDIES": "Summarize supporting studies that are part of this clinical study.",
                "ASSAYS/METHODOLOGIES": "Briefly describe any specialized assays or methodologies that will be used in this clinical study.",
                "STATISTICAL ANALYSIS PLAN": "Summarize the Statistical Analysis Plan or describe how the data will be analyzed.",
                "OUTCOME CRITERIA": "Describe criteria that would define whether you would or would not move forward with the subsequent development plan.",
                "RISKS": "Identify potential risks and mitigation strategies.",
                "CLINICAL SITES": "Indicate the number of clinical sites that will participate in the study.",
                "CLINICAL OPERATIONS PLAN": "Summarize the plan for managing the conduct of the clinical study.",
                "ENROLLMENT": "Describe the enrollment strategy and provide a timeline showing enrollment projections. Describe plans for inclusion of women and minorities.",
                "LONG TERM FOLLOW UP": "Describe requirements and plans for long term follow up and indicate how these will be supported.",
                "TIMELINE": "Provide a timeline for completion of the study and indicate relevant milestones."
            }
            
            # Extract fields with progress tracking
            extracted_fields = {}
            progress_bar = st.progress(0)
            total_fields = len(field_definitions)
            
            for i, (field_title, field_desc) in enumerate(field_definitions.items()):
                with st.spinner(f"📝 Analyzing {field_title}..."):
                    extracted_fields[field_title] = extract_field_with_claude(
                        parsed_content, field_title, field_desc, anthropic_client
                    )
                progress_bar.progress((i + 1) / total_fields)
            
            st.success("✅ Document processed successfully!")
            
            # Display results in an expandable section
            with st.expander("View Extracted Information", expanded=True):
                for field_title, content in extracted_fields.items():
                    st.subheader(field_title)
                    st.write(content)
            
            # Generate and provide download for DOCX
            with st.spinner("📑 Generating DOCX summary..."):
                doc_path = generate_docx_summary(extracted_fields)
            
            with open(doc_path, "rb") as file:
                st.download_button(
                    "📄 Download Summary (DOCX)",
                    file,
                    file_name="protocol_summary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        except Exception as e:
            st.error(f"❌ An error occurred: {str(e)}")

if __name__ == "__main__":
    main()
